# encoding: utf-8
"""
Author: gugan1
Info：pdf转换为excel、word、txt
CreateTime: 2022年11月9日
UpdateTime: 2022年11月10日
"""

import pdfplumber  # pip --default-timeout=1688 install pdfplumber -i http://pypi.douban.com/simple/ --trusted-host pypi.douban.com
from openpyxl import Workbook  #pip install openpyxl 
import os

import docx # pip install python-docx
from docx import Document


# file_path=os.getcwd()+"/2.pdf"
#pdf的文件地址
def Pdf_to_excel(pdf, file_path):
    wb = Workbook()  # 创建文件对象
    ws = wb.active  # 获取第一个sheet
    page_tag = 1
    for page in pdf.pages:
        # 获取当前页面的全部文本信息，包括表格中的文字
        #print(page.extract_text())
        for table in page.extract_tables():
            # print(table)
            for row in table:
                # print(row)
                #把列表拆了
                rowlist = str(row).replace(
                    "[",
                    "",
                ).replace("]", "").replace("'", "").replace("\\n",
                                                            "").split(",")
                #print(rowlist)
                ws.append(rowlist)
        print('------处理第%d页------\n' % page_tag)
        page_tag += 1
    pdf.close()
    # end_file = file_path.replace('pdf', 'xlsx').split("\\")[-1]
    end_file = file_path.replace('.pdf', '.xlsx')
    wb.save(end_file)
    print('写入excel成功')
    print('保存位置：%s'%end_file)




def Pdf_to_word(pdf,file_path):
    doc = docx.Document()   
    paragraph = doc.add_paragraph()
    page_tag = 1
    for page in pdf.pages:
        print('------处理第%d页------\n' % page_tag)
        text = page.extract_text()
        paragraph.add_run(text)
        page_tag+=1
        
    pdf.close()
    end_file = file_path.replace('.pdf', '.docx')
    doc.save(end_file) 
    print('写入word成功')
    print('保存位置：%s'%end_file)


def Pdf_to_txt(pdf,file_path):
    end_file=file_path.replace(".pdf", ".txt")
    with open(end_file,'a',encoding='utf-8') as file:
        page_tag = 1
        for page in pdf.pages:
            print('------处理第%d页------\n' % page_tag)
            text=page.extract_text()
            file.write(text)
            page_tag+=1
        pdf.close()
        print('写入txt成功')
        print('保存位置：%s'%end_file)
         
    


if __name__ == '__main__':
    
    choose=input("请选择转换格式，默认为excel直接点击回车\n输入1转换word\n输入2转换为txt\n")
    if not choose:
        path = input("请将整个文件或者文件夹拖入窗口\n")
        if path.endswith(".pdf"):
            ######单个文件操作
            file_path = path
            try:
                pdf = pdfplumber.open(file_path)  #打开pdf文件
                print('开始读取数据\n')
                Pdf_to_excel(pdf, file_path)
            except:
                print("操作错误！")
        else:
        ######批量转换文件夹中的pdf文件
            try:
                file_names = os.listdir(path)
                for file_name in file_names:
                    if file_name.endswith(".pdf"):
                        file_path = path + '\\' + file_name
                        pdf = pdfplumber.open(file_path)  #打开pdf文件
                        print('开始读取数据--%s' % file_name)
                        Pdf_to_excel(pdf, file_path)
            except:
                print("操作错误！")
    elif choose=='1':
        #转换成word格式
        path = input("请将整个文件或者文件夹拖入窗口\n")
        if path.endswith(".pdf"):
            ######单个文件操作
            file_path = path
            try:
                pdf = pdfplumber.open(file_path)  #打开pdf文件
                print('开始读取数据\n')
                Pdf_to_word(pdf, file_path)
            except:
                print("操作错误！")
        else:
        ######批量转换文件夹中的pdf文件
            try:
                file_names = os.listdir(path)
                for file_name in file_names:
                    if file_name.endswith(".pdf"):
                        file_path = path + '\\' + file_name
                        pdf = pdfplumber.open(file_path)  #打开pdf文件
                        print('开始读取数据--%s' % file_name)
                        Pdf_to_word(pdf, file_path)
            except:
                print("操作错误！")
    elif choose=='2':
        path = input("请将整个文件或者文件夹拖入窗口\n")
        if path.endswith(".pdf"):
            ######单个文件操作
            file_path = path
            try:
                pdf = pdfplumber.open(file_path)  #打开pdf文件
                print('开始读取数据\n')
                Pdf_to_txt(pdf, file_path)
            except:
                print("操作错误！")
        else:
        ######批量转换文件夹中的pdf文件
            try:
                file_names = os.listdir(path)
                for file_name in file_names:
                    if file_name.endswith(".pdf"):
                        file_path = path + '\\' + file_name
                        pdf = pdfplumber.open(file_path)  #打开pdf文件
                        print('开始读取数据--%s' % file_name)
                        Pdf_to_txt(pdf, file_path)
            except:
                print("操作错误！")
    else:
        print("输入错误")

            